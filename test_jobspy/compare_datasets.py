# -*- coding: utf-8 -*-
"""
Compare two job datasets side by side
Creates a detailed Word document comparing salary, job counts, industry distribution, etc.
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import pandas as pd
import numpy as np
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from pathlib import Path
import re

# Set matplotlib to use a backend that works without display
plt.switch_backend('Agg')
# Set style
sns.set_style("whitegrid")
plt.rcParams['font.sans-serif'] = ['SimHei', 'Arial Unicode MS', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

# ============================================================================
# COMPARISON CONFIGURATION - Modify these settings as needed
# ============================================================================

# Paths to the two datasets to compare
# Examples:
#   DATASET1_PATH = "output/BunchGlobal_2025_12_26"
#   DATASET2_PATH = "output/BunchGlobal_2025_01_18"
DATASET1_PATH = "output/BunchGlobal_2025_12_26"
DATASET2_PATH = "output/BunchGlobal_2025_01_18"
DATASET1_LABEL = "December 2025"
DATASET2_LABEL = "January 2025"

# Excel file name to look for in each region subdirectory
EXCEL_FILENAME = "jobspy_max_output.xlsx"

# Region name mapping (folder name -> display name)
REGION_NAME_MAP = {
    "united_states": "United States",
    "united_kingdom": "United Kingdom",
    "hong_kong": "Hong Kong",
    "singapore": "Singapore",
    "australia": "Australia",
}

# ============================================================================
# END OF CONFIGURATION
# ============================================================================


def map_excel_columns_to_standard(df):
    """Map Excel column names to standard snake_case names"""
    column_mapping = {
        "Job Title": "job_title",
        "Company Name": "company_name",
        "Requirements": "requirements",
        "Location": "location",
        "Salary Range": "salary_range",
        "Estimated Annual Salary": "estimated_annual_salary",
        "Estimated Annual Salary (USD)": "estimated_annual_salary_usd",
        "Job Description": "job_description",
        "Team Size/Business Line Size": "team_size",
        "Company Size": "company_size",
        "Posted Date": "posted_date",
        "Job Status": "job_status",
        "Platform": "platform",
        "Job Link": "job_link",
    }
    return df.rename(columns=column_mapping)


def load_data_from_folder(folder_path):
    """Load all Excel files from a folder structure"""
    folder_path = Path(folder_path)
    
    if not folder_path.exists():
        raise ValueError(f"Folder does not exist: {folder_path}")
    
    all_data = {}
    
    subdirs = [d for d in folder_path.iterdir() if d.is_dir()]
    
    for subdir in subdirs:
        excel_file = subdir / EXCEL_FILENAME
        
        if not excel_file.exists():
            continue
        
        region_name = REGION_NAME_MAP.get(subdir.name, subdir.name.replace('_', ' ').title())
        
        try:
            df = pd.read_excel(excel_file)
            if df.empty:
                all_data[region_name] = pd.DataFrame()
                continue
            
            df = map_excel_columns_to_standard(df)
            all_data[region_name] = df
            
        except Exception as e:
            print(f"    Error loading {excel_file}: {str(e)[:100]}")
            all_data[region_name] = pd.DataFrame()
    
    return all_data


def extract_salary_value(salary_str):
    """Extract numeric salary value from string"""
    if pd.isna(salary_str) or not salary_str:
        return None
    
    salary_str = str(salary_str).strip()
    if not salary_str or salary_str == 'nan':
        return None
    
    patterns = [
        r'[\$£€A\$S\$HK\$C\$]?\s*([\d,]+)\s*[kK]',
        r'[\$£€A\$S\$HK\$C\$]?\s*([\d,]+)',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, salary_str)
        if match:
            value_str = match.group(1).replace(',', '')
            try:
                value = float(value_str)
                if 'k' in salary_str.lower():
                    value *= 1000
                return value
            except:
                pass
    
    return None


def analyze_salary(df):
    """Analyze salary data"""
    if df.empty:
        return {'has_data': False}
    
    usd_salaries = []
    for salary_str in df.get('estimated_annual_salary_usd', []):
        salary_val = extract_salary_value(salary_str)
        if salary_val and 10000 <= salary_val <= 500000:
            usd_salaries.append(salary_val)
    
    if not usd_salaries:
        return {'has_data': False}
    
    usd_salaries = np.array(usd_salaries)
    
    return {
        'has_data': True,
        'count': len(usd_salaries),
        'mean': float(np.mean(usd_salaries)),
        'median': float(np.median(usd_salaries)),
        'std': float(np.std(usd_salaries)),
        'min': float(np.min(usd_salaries)),
        'max': float(np.max(usd_salaries)),
        'q25': float(np.percentile(usd_salaries, 25)),
        'q75': float(np.percentile(usd_salaries, 75)),
        'values': usd_salaries.tolist()
    }


def categorize_job_title(title):
    """Categorize job title into industry/role category"""
    if pd.isna(title) or not title:
        return "Other"
    
    title_lower = str(title).lower()
    
    # AI/ML categories
    if any(kw in title_lower for kw in ['ai engineer', 'machine learning', 'ml engineer', 'deep learning']):
        return "AI/ML Engineer"
    elif any(kw in title_lower for kw in ['data scientist', 'data science', 'data analyst']):
        return "Data Scientist/Analyst"
    elif any(kw in title_lower for kw in ['nlp', 'natural language', 'computer vision', 'cv engineer']):
        return "AI Specialist (NLP/CV)"
    elif any(kw in title_lower for kw in ['ai product', 'ai pm', 'ai manager']):
        return "AI Product/Management"
    elif any(kw in title_lower for kw in ['ai architect', 'ai solution', 'ai platform']):
        return "AI Architecture"
    elif any(kw in title_lower for kw in ['ai researcher', 'ml researcher', 'research scientist']):
        return "AI Research"
    elif any(kw in title_lower for kw in ['ai', 'artificial intelligence', 'ml', 'machine learning']):
        return "AI/ML (Other)"
    
    # Software engineering
    elif any(kw in title_lower for kw in ['software engineer', 'software developer', 'backend engineer', 'frontend engineer']):
        return "Software Engineer"
    elif any(kw in title_lower for kw in ['senior engineer', 'lead engineer', 'principal engineer']):
        return "Senior/Lead Engineer"
    
    # Other tech roles
    elif any(kw in title_lower for kw in ['product manager', 'pm', 'product owner']):
        return "Product Manager"
    elif any(kw in title_lower for kw in ['data engineer', 'etl', 'data pipeline']):
        return "Data Engineer"
    elif any(kw in title_lower for kw in ['devops', 'sre', 'site reliability']):
        return "DevOps/SRE"
    elif any(kw in title_lower for kw in ['qa', 'quality assurance', 'test engineer']):
        return "QA/Testing"
    
    return "Other"


def analyze_job_categories(df):
    """Analyze job title categories"""
    if df.empty or 'job_title' not in df.columns:
        return {'has_data': False}
    
    df_with_categories = df.copy()
    df_with_categories['category'] = df_with_categories['job_title'].apply(categorize_job_title)
    
    category_counts = df_with_categories['category'].value_counts()
    category_percentages = (category_counts / len(df_with_categories) * 100).round(2)
    
    return {
        'has_data': True,
        'counts': category_counts.to_dict(),
        'percentages': category_percentages.to_dict(),
        'total': len(df_with_categories)
    }


def analyze_platform_distribution(df):
    """Analyze platform distribution"""
    if df.empty or 'platform' not in df.columns:
        return {'has_data': False}
    
    platforms = df['platform'].dropna()
    platforms = platforms[platforms.astype(str).str.strip() != '']
    
    if platforms.empty:
        return {'has_data': False}
    
    platform_counts = platforms.value_counts()
    platform_percentages = (platform_counts / len(platforms) * 100).round(2)
    
    return {
        'has_data': True,
        'counts': platform_counts.to_dict(),
        'percentages': platform_percentages.to_dict(),
        'total': len(platforms)
    }


def analyze_location_distribution(df):
    """Analyze top locations"""
    if df.empty or 'location' not in df.columns:
        return {'has_data': False}
    
    locations = df['location'].dropna()
    locations = locations[locations.astype(str).str.strip() != '']
    
    if locations.empty:
        return {'has_data': False}
    
    location_counts = locations.value_counts()
    
    return {
        'has_data': True,
        'top_locations': location_counts.head(10).to_dict(),
        'total_unique': len(location_counts)
    }


def create_comparison_chart_salary(data1, data2, label1, label2, output_dir):
    """Create salary comparison chart"""
    salary_data = {}
    
    for region_name in set(list(data1.keys()) + list(data2.keys())):
        salaries1 = []
        salaries2 = []
        
        if region_name in data1 and not data1[region_name].empty:
            salary_info1 = analyze_salary(data1[region_name])
            if salary_info1.get('has_data'):
                salaries1 = salary_info1['values']
        
        if region_name in data2 and not data2[region_name].empty:
            salary_info2 = analyze_salary(data2[region_name])
            if salary_info2.get('has_data'):
                salaries2 = salary_info2['values']
        
        if salaries1 or salaries2:
            salary_data[region_name] = {
                label1: salaries1,
                label2: salaries2
            }
    
    if not salary_data:
        return None
    
    fig, axes = plt.subplots(2, 1, figsize=(14, 10))
    
    # Box plot comparison
    ax1 = axes[0]
    regions = list(salary_data.keys())
    positions = np.arange(len(regions))
    width = 0.35
    
    data1_for_box = [salary_data[r][label1] for r in regions]
    data2_for_box = [salary_data[r][label2] for r in regions]
    
    bp1 = ax1.boxplot(data1_for_box, positions=positions - width/2, widths=width, 
                      patch_artist=True)
    bp2 = ax1.boxplot(data2_for_box, positions=positions + width/2, widths=width,
                      patch_artist=True)
    ax1.set_xticks(positions)
    ax1.set_xticklabels(regions)
    
    for patch in bp1['boxes']:
        patch.set_facecolor('lightblue')
    for patch in bp2['boxes']:
        patch.set_facecolor('lightcoral')
    
    ax1.set_ylabel('Annual Salary (USD)', fontsize=12)
    ax1.set_title(f'Salary Distribution Comparison: {label1} vs {label2}', fontsize=14, fontweight='bold')
    ax1.legend([bp1['boxes'][0], bp2['boxes'][0]], [label1, label2])
    ax1.grid(True, alpha=0.3)
    plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, ha='right')
    
    # Mean comparison bar chart
    ax2 = axes[1]
    means1 = [np.mean(salary_data[r][label1]) if salary_data[r][label1] else 0 for r in regions]
    means2 = [np.mean(salary_data[r][label2]) if salary_data[r][label2] else 0 for r in regions]
    
    x = np.arange(len(regions))
    bars1 = ax2.bar(x - width/2, means1, width, label=label1, color='lightblue')
    bars2 = ax2.bar(x + width/2, means2, width, label=label2, color='lightcoral')
    
    ax2.set_ylabel('Mean Annual Salary (USD)', fontsize=12)
    ax2.set_xlabel('Region', fontsize=12)
    ax2.set_title('Mean Salary Comparison by Region', fontsize=14, fontweight='bold')
    ax2.set_xticks(x)
    ax2.set_xticklabels(regions, rotation=45, ha='right')
    ax2.legend()
    ax2.grid(True, alpha=0.3, axis='y')
    
    plt.tight_layout()
    
    chart_path = os.path.join(output_dir, 'salary_comparison_chart.png')
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return chart_path


def create_comparison_chart_job_counts(data1, data2, label1, label2, output_dir):
    """Create job count comparison chart"""
    regions = sorted(set(list(data1.keys()) + list(data2.keys())))
    
    counts1 = [len(data1.get(r, pd.DataFrame())) for r in regions]
    counts2 = [len(data2.get(r, pd.DataFrame())) for r in regions]
    
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))
    
    # Bar chart
    x = np.arange(len(regions))
    width = 0.35
    
    bars1 = ax1.bar(x - width/2, counts1, width, label=label1, color='lightblue')
    bars2 = ax1.bar(x + width/2, counts2, width, label=label2, color='lightcoral')
    
    ax1.set_ylabel('Number of Jobs', fontsize=12)
    ax1.set_xlabel('Region', fontsize=12)
    ax1.set_title('Job Count Comparison by Region', fontsize=14, fontweight='bold')
    ax1.set_xticks(x)
    ax1.set_xticklabels(regions, rotation=45, ha='right')
    ax1.legend()
    ax1.grid(True, alpha=0.3, axis='y')
    
    # Add value labels
    for bars in [bars1, bars2]:
        for bar in bars:
            height = bar.get_height()
            if height > 0:
                ax1.text(bar.get_x() + bar.get_width()/2., height,
                        f'{int(height)}',
                        ha='center', va='bottom', fontsize=9)
    
    # Percentage change
    changes = []
    change_labels = []
    for i, region in enumerate(regions):
        if counts1[i] > 0:
            change = ((counts2[i] - counts1[i]) / counts1[i] * 100)
            changes.append(change)
            change_labels.append(region)
    
    if changes:
        colors = ['green' if c > 0 else 'red' for c in changes]
        ax2.barh(change_labels, changes, color=colors, alpha=0.7)
        ax2.axvline(x=0, color='black', linestyle='--', linewidth=1)
        ax2.set_xlabel('Percentage Change (%)', fontsize=12)
        ax2.set_title('Job Count Change (%)', fontsize=14, fontweight='bold')
        ax2.grid(True, alpha=0.3, axis='x')
        
        for i, (label, change) in enumerate(zip(change_labels, changes)):
            ax2.text(change, i, f'{change:+.1f}%', 
                    va='center', ha='left' if change > 0 else 'right', fontsize=9)
    
    plt.tight_layout()
    
    chart_path = os.path.join(output_dir, 'job_count_comparison_chart.png')
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return chart_path


def create_comparison_chart_categories(data1, data2, label1, label2, output_dir):
    """Create job category comparison chart"""
    # Aggregate categories across all regions
    all_categories1 = {}
    all_categories2 = {}
    
    for region_name in set(list(data1.keys()) + list(data2.keys())):
        if region_name in data1 and not data1[region_name].empty:
            cat_info1 = analyze_job_categories(data1[region_name])
            if cat_info1.get('has_data'):
                for cat, count in cat_info1['counts'].items():
                    all_categories1[cat] = all_categories1.get(cat, 0) + count
        
        if region_name in data2 and not data2[region_name].empty:
            cat_info2 = analyze_job_categories(data2[region_name])
            if cat_info2.get('has_data'):
                for cat, count in cat_info2['counts'].items():
                    all_categories2[cat] = all_categories2.get(cat, 0) + count
    
    if not all_categories1 and not all_categories2:
        return None
    
    all_categories = sorted(set(list(all_categories1.keys()) + list(all_categories2.keys())))
    
    counts1 = [all_categories1.get(cat, 0) for cat in all_categories]
    counts2 = [all_categories2.get(cat, 0) for cat in all_categories]
    
    # Calculate percentages
    total1 = sum(counts1) if sum(counts1) > 0 else 1
    total2 = sum(counts2) if sum(counts2) > 0 else 1
    percentages1 = [c / total1 * 100 for c in counts1]
    percentages2 = [c / total2 * 100 for c in counts2]
    
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 8))
    
    # Count comparison
    x = np.arange(len(all_categories))
    width = 0.35
    
    bars1 = ax1.bar(x - width/2, counts1, width, label=label1, color='lightblue')
    bars2 = ax1.bar(x + width/2, counts2, width, label=label2, color='lightcoral')
    
    ax1.set_ylabel('Number of Jobs', fontsize=12)
    ax1.set_xlabel('Job Category', fontsize=12)
    ax1.set_title('Job Category Distribution (Counts)', fontsize=14, fontweight='bold')
    ax1.set_xticks(x)
    ax1.set_xticklabels(all_categories, rotation=45, ha='right')
    ax1.legend()
    ax1.grid(True, alpha=0.3, axis='y')
    
    # Percentage comparison
    bars3 = ax2.bar(x - width/2, percentages1, width, label=label1, color='lightblue')
    bars4 = ax2.bar(x + width/2, percentages2, width, label=label2, color='lightcoral')
    
    ax2.set_ylabel('Percentage (%)', fontsize=12)
    ax2.set_xlabel('Job Category', fontsize=12)
    ax2.set_title('Job Category Distribution (Percentages)', fontsize=14, fontweight='bold')
    ax2.set_xticks(x)
    ax2.set_xticklabels(all_categories, rotation=45, ha='right')
    ax2.legend()
    ax2.grid(True, alpha=0.3, axis='y')
    
    plt.tight_layout()
    
    chart_path = os.path.join(output_dir, 'category_comparison_chart.png')
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return chart_path


def add_heading(doc, text, level=1):
    """Add a heading to document"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph to document"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_comparison_table(doc, data1, data2, label1, label2, headers):
    """Add a comparison table to document"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    for key in sorted(set(list(data1.keys()) + list(data2.keys()))):
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(data1.get(key, 'N/A'))
        row_cells[2].text = str(data2.get(key, 'N/A'))
        
        # Calculate change if both values are numeric
        try:
            val1 = float(data1.get(key, 0)) if data1.get(key, 'N/A') != 'N/A' else 0
            val2 = float(data2.get(key, 0)) if data2.get(key, 'N/A') != 'N/A' else 0
            if val1 > 0:
                change = ((val2 - val1) / val1 * 100)
                row_cells[3].text = f"{change:+.1f}%"
            else:
                row_cells[3].text = "N/A"
        except:
            row_cells[3].text = "N/A"
    
    return table


def generate_comparison_report(data1, data2, label1, label2, output_path, path1, path2):
    """Generate comprehensive comparison Word report"""
    print(f"\nGenerating comparison report: {output_path}")
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Job Data Comparison Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Report metadata
    p = doc.add_paragraph()
    p.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run(f'Dataset 1: {label1} ({path1})').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run(f'Dataset 2: {label2} ({path2})').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Executive Summary
    add_heading(doc, 'Executive Summary', 1)
    
    total_jobs1 = sum(len(df) for df in data1.values() if not df.empty)
    total_jobs2 = sum(len(df) for df in data2.values() if not df.empty)
    job_change = ((total_jobs2 - total_jobs1) / total_jobs1 * 100) if total_jobs1 > 0 else 0
    
    summary_text = f"""
This report compares two job datasets collected approximately one month apart.
Dataset 1 ({label1}) contains {total_jobs1:,} total job records.
Dataset 2 ({label2}) contains {total_jobs2:,} total job records.
Overall job count change: {job_change:+.1f}%

The comparison covers:
- Total job counts by region
- Salary distributions and trends
- Job category/industry distribution
- Platform distribution (Indeed vs LinkedIn)
- Geographic distribution
    """
    add_paragraph(doc, summary_text.strip())
    
    doc.add_page_break()
    
    # Overall Statistics Comparison
    add_heading(doc, 'Overall Statistics Comparison', 1)
    
    overall_stats = {
        'Total Jobs': (total_jobs1, total_jobs2),
        'Regions with Data': (len([r for r, df in data1.items() if not df.empty]),
                            len([r for r, df in data2.items() if not df.empty]))
    }
    
    stats_table = doc.add_table(rows=1, cols=4)
    stats_table.style = 'Light Grid Accent 1'
    header_cells = stats_table.rows[0].cells
    headers = ['Metric', label1, label2, 'Change']
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for key, (val1, val2) in overall_stats.items():
        row_cells = stats_table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(val1)
        row_cells[2].text = str(val2)
        if val1 > 0:
            change = ((val2 - val1) / val1 * 100)
            row_cells[3].text = f"{change:+.1f}%"
        else:
            row_cells[3].text = "N/A"
    
    doc.add_page_break()
    
    # Job Count Comparison by Region
    add_heading(doc, 'Job Count Comparison by Region', 1)
    
    region_counts1 = {r: len(df) for r, df in data1.items() if not df.empty}
    region_counts2 = {r: len(df) for r, df in data2.items() if not df.empty}
    
    add_comparison_table(doc, region_counts1, region_counts2, label1, label2,
                        ['Region', label1, label2, 'Change (%)'])
    
    doc.add_page_break()
    
    # Salary Comparison
    add_heading(doc, 'Salary Comparison', 1)
    
    # Aggregate salary statistics
    all_salaries1 = []
    all_salaries2 = []
    
    for region_name in set(list(data1.keys()) + list(data2.keys())):
        if region_name in data1 and not data1[region_name].empty:
            salary_info1 = analyze_salary(data1[region_name])
            if salary_info1.get('has_data'):
                all_salaries1.extend(salary_info1['values'])
        
        if region_name in data2 and not data2[region_name].empty:
            salary_info2 = analyze_salary(data2[region_name])
            if salary_info2.get('has_data'):
                all_salaries2.extend(salary_info2['values'])
    
    if all_salaries1 or all_salaries2:
        add_heading(doc, 'Overall Salary Statistics', 2)
        
        salary_stats = {}
        if all_salaries1:
            salary_stats['Mean Salary (USD)'] = (np.mean(all_salaries1), np.mean(all_salaries2) if all_salaries2 else 0)
            salary_stats['Median Salary (USD)'] = (np.median(all_salaries1), np.median(all_salaries2) if all_salaries2 else 0)
            salary_stats['25th Percentile'] = (np.percentile(all_salaries1, 25), np.percentile(all_salaries2, 25) if all_salaries2 else 0)
            salary_stats['75th Percentile'] = (np.percentile(all_salaries1, 75), np.percentile(all_salaries2, 75) if all_salaries2 else 0)
        
        salary_table = doc.add_table(rows=1, cols=4)
        salary_table.style = 'Light Grid Accent 1'
        header_cells = salary_table.rows[0].cells
        for i, header in enumerate(['Metric', label1, label2, 'Change (%)']):
            header_cells[i].text = str(header)
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        for key, (val1, val2) in salary_stats.items():
            row_cells = salary_table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = f"${val1:,.0f}"
            row_cells[2].text = f"${val2:,.0f}" if val2 > 0 else "N/A"
            if val1 > 0 and val2 > 0:
                change = ((val2 - val1) / val1 * 100)
                row_cells[3].text = f"{change:+.1f}%"
            else:
                row_cells[3].text = "N/A"
        
        # Regional salary comparison
        add_heading(doc, 'Salary Statistics by Region', 2)
        
        for region_name in sorted(set(list(data1.keys()) + list(data2.keys()))):
            salary_info1 = analyze_salary(data1.get(region_name, pd.DataFrame()))
            salary_info2 = analyze_salary(data2.get(region_name, pd.DataFrame()))
            
            if salary_info1.get('has_data') or salary_info2.get('has_data'):
                add_heading(doc, region_name, 3)
                
                region_salary_table = doc.add_table(rows=1, cols=4)
                region_salary_table.style = 'Light Grid Accent 1'
                header_cells = region_salary_table.rows[0].cells
                for i, header in enumerate(['Metric', label1, label2, 'Change (%)']):
                    header_cells[i].text = str(header)
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
                
                metrics = ['Mean', 'Median', '25th Percentile', '75th Percentile']
                metric_keys = ['mean', 'median', 'q25', 'q75']
                
                for metric, key in zip(metrics, metric_keys):
                    row_cells = region_salary_table.add_row().cells
                    row_cells[0].text = metric
                    val1 = salary_info1.get(key, 0) if salary_info1.get('has_data') else 0
                    val2 = salary_info2.get(key, 0) if salary_info2.get('has_data') else 0
                    row_cells[1].text = f"${val1:,.0f}" if val1 > 0 else "N/A"
                    row_cells[2].text = f"${val2:,.0f}" if val2 > 0 else "N/A"
                    if val1 > 0 and val2 > 0:
                        change = ((val2 - val1) / val1 * 100)
                        row_cells[3].text = f"{change:+.1f}%"
                    else:
                        row_cells[3].text = "N/A"
    
    doc.add_page_break()
    
    # Job Category Comparison
    add_heading(doc, 'Job Category/Industry Comparison', 1)
    
    # Aggregate categories
    all_categories1 = {}
    all_categories2 = {}
    
    for region_name in set(list(data1.keys()) + list(data2.keys())):
        if region_name in data1 and not data1[region_name].empty:
            cat_info1 = analyze_job_categories(data1[region_name])
            if cat_info1.get('has_data'):
                for cat, count in cat_info1['counts'].items():
                    all_categories1[cat] = all_categories1.get(cat, 0) + count
        
        if region_name in data2 and not data2[region_name].empty:
            cat_info2 = analyze_job_categories(data2[region_name])
            if cat_info2.get('has_data'):
                for cat, count in cat_info2['counts'].items():
                    all_categories2[cat] = all_categories2.get(cat, 0) + count
    
    if all_categories1 or all_categories2:
        add_heading(doc, 'Category Distribution (Counts)', 2)
        add_comparison_table(doc, all_categories1, all_categories2, label1, label2,
                            ['Category', label1, label2, 'Change (%)'])
        
        # Calculate percentages
        total1 = sum(all_categories1.values()) if all_categories1 else 1
        total2 = sum(all_categories2.values()) if all_categories2 else 1
        percentages1 = {k: v / total1 * 100 for k, v in all_categories1.items()}
        percentages2 = {k: v / total2 * 100 for k, v in all_categories2.items()}
        
        add_heading(doc, 'Category Distribution (Percentages)', 2)
        add_comparison_table(doc, percentages1, percentages2, label1, label2,
                            ['Category', f'{label1} (%)', f'{label2} (%)', 'Change (pp)'])
    
    doc.add_page_break()
    
    # Platform Comparison
    add_heading(doc, 'Platform Distribution Comparison', 1)
    
    all_platforms1 = {}
    all_platforms2 = {}
    
    for region_name in set(list(data1.keys()) + list(data2.keys())):
        if region_name in data1 and not data1[region_name].empty:
            platform_info1 = analyze_platform_distribution(data1[region_name])
            if platform_info1.get('has_data'):
                for platform, count in platform_info1['counts'].items():
                    all_platforms1[platform] = all_platforms1.get(platform, 0) + count
        
        if region_name in data2 and not data2[region_name].empty:
            platform_info2 = analyze_platform_distribution(data2[region_name])
            if platform_info2.get('has_data'):
                for platform, count in platform_info2['counts'].items():
                    all_platforms2[platform] = all_platforms2.get(platform, 0) + count
    
    if all_platforms1 or all_platforms2:
        add_comparison_table(doc, all_platforms1, all_platforms2, label1, label2,
                            ['Platform', label1, label2, 'Change (%)'])
    
    doc.add_page_break()
    
    # Charts Section
    add_heading(doc, 'Visual Comparisons', 1)
    
    charts_dir = os.path.join(os.path.dirname(output_path), 'comparison_charts')
    os.makedirs(charts_dir, exist_ok=True)
    
    # Salary comparison chart
    chart_path = create_comparison_chart_salary(data1, data2, label1, label2, charts_dir)
    if chart_path:
        add_heading(doc, 'Salary Distribution Comparison', 2)
        add_paragraph(doc, 'Salary distributions across regions, comparing both datasets.')
        doc.add_picture(chart_path, width=Inches(6))
    
    # Job count comparison chart
    chart_path = create_comparison_chart_job_counts(data1, data2, label1, label2, charts_dir)
    if chart_path:
        add_heading(doc, 'Job Count Comparison', 2)
        add_paragraph(doc, 'Total job counts by region and percentage changes.')
        doc.add_picture(chart_path, width=Inches(6))
    
    # Category comparison chart
    chart_path = create_comparison_chart_categories(data1, data2, label1, label2, charts_dir)
    if chart_path:
        add_heading(doc, 'Job Category Distribution Comparison', 2)
        add_paragraph(doc, 'Distribution of jobs across different categories/industries.')
        doc.add_picture(chart_path, width=Inches(6))
    
    # Save document
    doc.save(output_path)
    print(f"Comparison report saved to: {output_path}")
    
    return output_path


def main():
    """Main function"""
    print("="*60)
    print("Job Data Comparison Report Generator")
    print("="*60)
    
    if not DATASET1_PATH or not DATASET2_PATH:
        print("[ERROR] Please configure DATASET1_PATH and DATASET2_PATH in the script")
        return
    
    try:
        # Load data from both folders
        print(f"\nLoading Dataset 1: {DATASET1_PATH}")
        data1 = load_data_from_folder(DATASET1_PATH)
        total1 = sum(len(df) for df in data1.values() if not df.empty)
        print(f"  Loaded {total1:,} total records")
        
        print(f"\nLoading Dataset 2: {DATASET2_PATH}")
        data2 = load_data_from_folder(DATASET2_PATH)
        total2 = sum(len(df) for df in data2.values() if not df.empty)
        print(f"  Loaded {total2:,} total records")
        
        if total1 == 0 and total2 == 0:
            print("\n[WARNING] No data found in either dataset. Cannot generate report.")
            return
        
        # Generate report
        output_dir = "output/reports"
        os.makedirs(output_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        folder1_name = Path(DATASET1_PATH).name
        folder2_name = Path(DATASET2_PATH).name
        output_path = os.path.join(output_dir, f"comparison_report_{folder1_name}_vs_{folder2_name}_{timestamp}.docx")
        
        generate_comparison_report(data1, data2, DATASET1_LABEL, DATASET2_LABEL, 
                                   output_path, DATASET1_PATH, DATASET2_PATH)
        
        print("\n" + "="*60)
        print("Comparison report generation completed!")
        print(f"Report saved to: {output_path}")
        print("="*60)
        
    except Exception as e:
        print(f"\n[ERROR] {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()

