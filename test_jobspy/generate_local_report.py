# -*- coding: utf-8 -*-
"""
Generate comprehensive report from local Excel files
Creates a detailed Word document with statistics, charts, and cross-region analysis
Supports analyzing data from a specified folder containing region subdirectories
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import pandas as pd
import numpy as np
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import argparse
from pathlib import Path
from io import BytesIO

# Set matplotlib to use a backend that works without display
plt.switch_backend('Agg')
# Set style
sns.set_style("whitegrid")
plt.rcParams['font.sans-serif'] = ['SimHei', 'Arial Unicode MS', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

# ============================================================================
# REPORT CONFIGURATION - Modify these settings as needed
# ============================================================================

# Default folder path (can be overridden by command line argument)
# Examples:
#   DEFAULT_FOLDER = "output/BunchGlobal_2025_01_18"
#   DEFAULT_FOLDER = None  # Must specify via command line
DEFAULT_FOLDER = "output/BunchGlobal_2025_01_18"

# Excel file name to look for in each region subdirectory
EXCEL_FILENAME = "jobspy_max_output.xlsx"

# Region name mapping (folder name -> display name)
# If a folder name is not in this map, it will be used as-is with title case
REGION_NAME_MAP = {
    "united_states": "United States",
    "united_kingdom": "United Kingdom",
    "hong_kong": "Hong Kong",
    "singapore": "Singapore",
    "australia": "Australia",
}

# ============================================================================
# END OF CONFIGURATION
# ============================================================================


def map_excel_columns_to_standard(df):
    """
    Map Excel column names (Title Case) to standard snake_case names
    for compatibility with analysis functions
    
    Args:
        df: DataFrame with Excel column names
    
    Returns:
        DataFrame with standardized column names
    """
    # Mapping from Excel column names to standard names
    column_mapping = {
        "Job Title": "job_title",
        "Company Name": "company_name",
        "Requirements": "requirements",
        "Location": "location",
        "Salary Range": "salary_range",
        "Estimated Annual Salary": "estimated_annual_salary",
        "Estimated Annual Salary (USD)": "estimated_annual_salary_usd",
        "Job Description": "job_description",
        "Team Size/Business Line Size": "team_size",
        "Company Size": "company_size",
        "Posted Date": "posted_date",
        "Job Status": "job_status",
        "Platform": "platform",
        "Job Link": "job_link",
    }
    
    # Create a copy to avoid modifying original
    df_mapped = df.copy()
    
    # Rename columns
    df_mapped = df_mapped.rename(columns=column_mapping)
    
    return df_mapped


def load_data_from_folder(folder_path):
    """
    Load all Excel files from a folder structure
    
    Expected structure:
    folder_path/
        region1/
            jobspy_max_output.xlsx
        region2/
            jobspy_max_output.xlsx
        ...
    
    Args:
        folder_path: Path to the folder containing region subdirectories
    
    Returns:
        Dictionary mapping region names to DataFrames
    """
    folder_path = Path(folder_path)
    
    if not folder_path.exists():
        raise ValueError(f"Folder does not exist: {folder_path}")
    
    if not folder_path.is_dir():
        raise ValueError(f"Path is not a directory: {folder_path}")
    
    all_data = {}
    
    print(f"\nScanning folder: {folder_path}")
    
    # Look for subdirectories
    subdirs = [d for d in folder_path.iterdir() if d.is_dir()]
    
    if not subdirs:
        raise ValueError(f"No subdirectories found in {folder_path}")
    
    for subdir in subdirs:
        excel_file = subdir / EXCEL_FILENAME
        
        if not excel_file.exists():
            print(f"  Warning: {EXCEL_FILENAME} not found in {subdir.name}, skipping")
            continue
        
        # Get region display name
        region_name = REGION_NAME_MAP.get(subdir.name, subdir.name.replace('_', ' ').title())
        
        try:
            print(f"  Loading data from {region_name}...")
            df = pd.read_excel(excel_file)
            
            if df.empty:
                print(f"    Warning: File is empty")
                all_data[region_name] = pd.DataFrame()
                continue
            
            # Map column names to standard format
            df = map_excel_columns_to_standard(df)
            
            all_data[region_name] = df
            print(f"    Loaded {len(df)} records")
            
        except Exception as e:
            print(f"    Error loading {excel_file}: {str(e)[:100]}")
            all_data[region_name] = pd.DataFrame()
    
    return all_data


def calculate_completeness(df):
    """Calculate data completeness for each field"""
    if df.empty:
        return {}
    
    total = len(df)
    completeness = {}
    
    fields = [
        'job_title', 'company_name', 'requirements', 'location',
        'salary_range', 'estimated_annual_salary', 'estimated_annual_salary_usd',
        'job_description', 'team_size', 'company_size',
        'posted_date', 'job_status', 'platform', 'job_link'
    ]
    
    for field in fields:
        if field in df.columns:
            non_empty = df[field].notna() & (df[field] != '') & (df[field].astype(str).str.strip() != '')
            count = non_empty.sum()
            completeness[field] = {
                'count': count,
                'percentage': (count / total * 100) if total > 0 else 0
            }
        else:
            completeness[field] = {'count': 0, 'percentage': 0}
    
    return completeness


def extract_salary_value(salary_str):
    """Extract numeric salary value from string"""
    if pd.isna(salary_str) or not salary_str:
        return None
    
    salary_str = str(salary_str).strip()
    if not salary_str or salary_str == 'nan':
        return None
    
    # Remove currency symbols and extract numbers
    import re
    # Match patterns like $100,000 or $100k or 100,000
    patterns = [
        r'[\$£€A\$S\$HK\$C\$]?\s*([\d,]+)\s*[kK]',  # $100k or 100k
        r'[\$£€A\$S\$HK\$C\$]?\s*([\d,]+)',  # $100,000 or 100,000
    ]
    
    for pattern in patterns:
        match = re.search(pattern, salary_str)
        if match:
            value_str = match.group(1).replace(',', '')
            try:
                value = float(value_str)
                # If it's a K notation, multiply by 1000
                if 'k' in salary_str.lower():
                    value *= 1000
                return value
            except:
                pass
    
    return None


def analyze_salary(df, region_name):
    """Analyze salary data"""
    if df.empty:
        return {}
    
    # Extract USD salaries
    usd_salaries = []
    for salary_str in df.get('estimated_annual_salary_usd', []):
        salary_val = extract_salary_value(salary_str)
        if salary_val and 10000 <= salary_val <= 500000:  # Reasonable range
            usd_salaries.append(salary_val)
    
    if not usd_salaries:
        return {'has_data': False}
    
    usd_salaries = np.array(usd_salaries)
    
    return {
        'has_data': True,
        'count': len(usd_salaries),
        'mean': float(np.mean(usd_salaries)),
        'median': float(np.median(usd_salaries)),
        'std': float(np.std(usd_salaries)),
        'min': float(np.min(usd_salaries)),
        'max': float(np.max(usd_salaries)),
        'q25': float(np.percentile(usd_salaries, 25)),
        'q75': float(np.percentile(usd_salaries, 75)),
        'values': usd_salaries.tolist()
    }


def analyze_time_range(df):
    """Analyze posted date range"""
    if df.empty or 'posted_date' not in df.columns:
        return {}
    
    dates = []
    for date_str in df['posted_date']:
        if pd.notna(date_str) and date_str:
            try:
                date_obj = pd.to_datetime(date_str)
                dates.append(date_obj)
            except:
                pass
    
    if not dates:
        return {'has_data': False}
    
    dates = pd.to_datetime(dates)
    
    return {
        'has_data': True,
        'count': len(dates),
        'earliest': dates.min(),
        'latest': dates.max(),
        'span_days': (dates.max() - dates.min()).days,
        'dates': dates
    }


def analyze_location(df):
    """Analyze location distribution"""
    if df.empty or 'location' not in df.columns:
        return {}
    
    locations = df['location'].dropna()
    locations = locations[locations.astype(str).str.strip() != '']
    
    if locations.empty:
        return {'has_data': False}
    
    location_counts = locations.value_counts()
    
    return {
        'has_data': True,
        'total_unique': len(location_counts),
        'top_locations': location_counts.head(10).to_dict()
    }


def analyze_platform(df):
    """Analyze platform distribution"""
    if df.empty or 'platform' not in df.columns:
        return {}
    
    platforms = df['platform'].dropna()
    platforms = platforms[platforms.astype(str).str.strip() != '']
    
    if platforms.empty:
        return {'has_data': False}
    
    platform_counts = platforms.value_counts()
    
    return {
        'has_data': True,
        'distribution': platform_counts.to_dict()
    }


def create_chart_completeness(all_data, output_dir):
    """Create completeness comparison chart"""
    regions = []
    completeness_scores = []
    
    for region_name, df in all_data.items():
        if not df.empty:
            comp = calculate_completeness(df)
            # Calculate overall completeness
            total_fields = len(comp)
            filled_fields = sum(v['count'] for v in comp.values())
            total_records = len(df)
            overall = (filled_fields / (total_fields * total_records) * 100) if total_records > 0 else 0
            regions.append(region_name)
            completeness_scores.append(overall)
    
    if not regions:
        return None
    
    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.bar(regions, completeness_scores, color=sns.color_palette("husl", len(regions)))
    ax.set_ylabel('Overall Completeness (%)', fontsize=12)
    ax.set_xlabel('Region', fontsize=12)
    ax.set_title('Data Completeness by Region', fontsize=14, fontweight='bold')
    ax.set_ylim(0, 100)
    
    # Add value labels on bars
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
                f'{height:.1f}%',
                ha='center', va='bottom', fontsize=10)
    
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    
    chart_path = os.path.join(output_dir, 'completeness_chart.png')
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return chart_path


def create_chart_salary_distribution(all_data, output_dir):
    """Create salary distribution comparison chart"""
    salary_data = {}
    
    for region_name, df in all_data.items():
        if not df.empty:
            salary_info = analyze_salary(df, region_name)
            if salary_info.get('has_data'):
                salary_data[region_name] = salary_info['values']
    
    if not salary_data:
        return None
    
    fig, axes = plt.subplots(2, 1, figsize=(12, 10))
    
    # Box plot
    ax1 = axes[0]
    data_for_box = [salary_data[r] for r in salary_data.keys()]
    labels = list(salary_data.keys())
    bp = ax1.boxplot(data_for_box, tick_labels=labels, patch_artist=True)
    for patch in bp['boxes']:
        patch.set_facecolor(sns.color_palette("husl", len(labels))[bp['boxes'].index(patch)])
    ax1.set_ylabel('Annual Salary (USD)', fontsize=12)
    ax1.set_title('Salary Distribution by Region (Box Plot)', fontsize=14, fontweight='bold')
    ax1.grid(True, alpha=0.3)
    plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, ha='right')
    
    # Bar chart with statistics
    ax2 = axes[1]
    means = [np.mean(salary_data[r]) for r in labels]
    medians = [np.median(salary_data[r]) for r in labels]
    x = np.arange(len(labels))
    width = 0.35
    
    bars1 = ax2.bar(x - width/2, means, width, label='Mean', color='skyblue')
    bars2 = ax2.bar(x + width/2, medians, width, label='Median', color='lightcoral')
    
    ax2.set_ylabel('Annual Salary (USD)', fontsize=12)
    ax2.set_xlabel('Region', fontsize=12)
    ax2.set_title('Mean vs Median Salary by Region', fontsize=14, fontweight='bold')
    ax2.set_xticks(x)
    ax2.set_xticklabels(labels, rotation=45, ha='right')
    ax2.legend()
    ax2.grid(True, alpha=0.3, axis='y')
    
    plt.tight_layout()
    
    chart_path = os.path.join(output_dir, 'salary_distribution_chart.png')
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return chart_path


def create_chart_time_series(all_data, output_dir):
    """Create time series chart of job postings"""
    time_data = {}
    
    for region_name, df in all_data.items():
        if not df.empty:
            time_info = analyze_time_range(df)
            if time_info.get('has_data'):
                dates = time_info['dates']
                # Count by month
                monthly_counts = dates.to_period('M').value_counts().sort_index()
                time_data[region_name] = monthly_counts
    
    if not time_data:
        return None
    
    fig, ax = plt.subplots(figsize=(14, 6))
    
    for region_name, monthly_counts in time_data.items():
        ax.plot(monthly_counts.index.astype(str), monthly_counts.values, 
                marker='o', label=region_name, linewidth=2, markersize=6)
    
    ax.set_xlabel('Month', fontsize=12)
    ax.set_ylabel('Number of Job Postings', fontsize=12)
    ax.set_title('Job Postings Over Time by Region', fontsize=14, fontweight='bold')
    ax.legend()
    ax.grid(True, alpha=0.3)
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    
    chart_path = os.path.join(output_dir, 'time_series_chart.png')
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return chart_path


def create_chart_platform_distribution(all_data, output_dir):
    """Create platform distribution chart"""
    platform_data = {}
    
    for region_name, df in all_data.items():
        if not df.empty:
            platform_info = analyze_platform(df)
            if platform_info.get('has_data'):
                platform_data[region_name] = platform_info['distribution']
    
    if not platform_data:
        return None
    
    # Aggregate across all regions
    total_platforms = {}
    for region_data in platform_data.values():
        for platform, count in region_data.items():
            total_platforms[platform] = total_platforms.get(platform, 0) + count
    
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))
    
    # Overall distribution
    platforms = list(total_platforms.keys())
    counts = list(total_platforms.values())
    colors = sns.color_palette("husl", len(platforms))
    ax1.pie(counts, labels=platforms, autopct='%1.1f%%', colors=colors, startangle=90)
    ax1.set_title('Overall Platform Distribution', fontsize=14, fontweight='bold')
    
    # By region
    regions = list(platform_data.keys())
    indeed_counts = [platform_data[r].get('Indeed', 0) for r in regions]
    linkedin_counts = [platform_data[r].get('LinkedIn', 0) for r in regions]
    
    x = np.arange(len(regions))
    width = 0.35
    ax2.bar(x - width/2, indeed_counts, width, label='Indeed', color='#2164f3')
    ax2.bar(x + width/2, linkedin_counts, width, label='LinkedIn', color='#0077b5')
    ax2.set_xlabel('Region', fontsize=12)
    ax2.set_ylabel('Number of Jobs', fontsize=12)
    ax2.set_title('Platform Distribution by Region', fontsize=14, fontweight='bold')
    ax2.set_xticks(x)
    ax2.set_xticklabels(regions, rotation=45, ha='right')
    ax2.legend()
    ax2.grid(True, alpha=0.3, axis='y')
    
    plt.tight_layout()
    
    chart_path = os.path.join(output_dir, 'platform_distribution_chart.png')
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return chart_path


def create_chart_region_comparison(all_data, output_dir):
    """Create comprehensive region comparison chart"""
    regions = []
    job_counts = []
    salary_means = []
    
    for region_name, df in all_data.items():
        if not df.empty:
            regions.append(region_name)
            job_counts.append(len(df))
            salary_info = analyze_salary(df, region_name)
            if salary_info.get('has_data'):
                salary_means.append(salary_info['mean'])
            else:
                salary_means.append(0)
    
    if not regions:
        return None
    
    fig, axes = plt.subplots(2, 2, figsize=(14, 10))
    
    # Job count
    ax1 = axes[0, 0]
    bars1 = ax1.bar(regions, job_counts, color=sns.color_palette("husl", len(regions)))
    ax1.set_ylabel('Number of Jobs', fontsize=12)
    ax1.set_title('Total Job Count by Region', fontsize=12, fontweight='bold')
    ax1.tick_params(axis='x', rotation=45)
    for bar in bars1:
        height = bar.get_height()
        ax1.text(bar.get_x() + bar.get_width()/2., height,
                f'{int(height)}',
                ha='center', va='bottom', fontsize=9)
    
    # Average salary
    ax2 = axes[0, 1]
    valid_salaries = [(r, s) for r, s in zip(regions, salary_means) if s > 0]
    if valid_salaries:
        regions_valid, salaries_valid = zip(*valid_salaries)
        bars2 = ax2.bar(regions_valid, salaries_valid, color='lightgreen')
        ax2.set_ylabel('Average Salary (USD)', fontsize=12)
        ax2.set_title('Average Salary by Region', fontsize=12, fontweight='bold')
        ax2.tick_params(axis='x', rotation=45)
        for bar in bars2:
            height = bar.get_height()
            ax2.text(bar.get_x() + bar.get_width()/2., height,
                    f'${int(height):,}',
                    ha='center', va='bottom', fontsize=9)
    
    # Completeness
    ax3 = axes[1, 0]
    completeness_scores = []
    for region_name in regions:
        df = all_data[region_name]
        if not df.empty:
            comp = calculate_completeness(df)
            total_fields = len(comp)
            filled_fields = sum(v['count'] for v in comp.values())
            total_records = len(df)
            overall = (filled_fields / (total_fields * total_records) * 100) if total_records > 0 else 0
            completeness_scores.append(overall)
        else:
            completeness_scores.append(0)
    
    bars3 = ax3.bar(regions, completeness_scores, color='orange')
    ax3.set_ylabel('Completeness (%)', fontsize=12)
    ax3.set_title('Data Completeness by Region', fontsize=12, fontweight='bold')
    ax3.set_ylim(0, 100)
    ax3.tick_params(axis='x', rotation=45)
    for bar in bars3:
        height = bar.get_height()
        ax3.text(bar.get_x() + bar.get_width()/2., height,
                f'{height:.1f}%',
                ha='center', va='bottom', fontsize=9)
    
    # Platform split
    ax4 = axes[1, 1]
    indeed_counts = []
    linkedin_counts = []
    for region_name in regions:
        df = all_data[region_name]
        if not df.empty and 'platform' in df.columns:
            platforms = df['platform'].value_counts()
            indeed_counts.append(platforms.get('Indeed', 0))
            linkedin_counts.append(platforms.get('LinkedIn', 0))
        else:
            indeed_counts.append(0)
            linkedin_counts.append(0)
    
    x = np.arange(len(regions))
    width = 0.35
    ax4.bar(x - width/2, indeed_counts, width, label='Indeed', color='#2164f3')
    ax4.bar(x + width/2, linkedin_counts, width, label='LinkedIn', color='#0077b5')
    ax4.set_ylabel('Number of Jobs', fontsize=12)
    ax4.set_title('Platform Split by Region', fontsize=12, fontweight='bold')
    ax4.set_xticks(x)
    ax4.set_xticklabels(regions, rotation=45, ha='right')
    ax4.legend()
    
    plt.tight_layout()
    
    chart_path = os.path.join(output_dir, 'region_comparison_chart.png')
    plt.savefig(chart_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return chart_path


def add_heading(doc, text, level=1):
    """Add a heading to document"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph to document"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_table_from_dict(doc, data_dict, headers=None):
    """Add a table from dictionary"""
    if not data_dict:
        return None
    
    if headers is None:
        headers = ['Field', 'Value']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    for key, value in data_dict.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)
    
    return table


def generate_report(all_data, output_path, source_folder=None):
    """
    Generate comprehensive Word report
    
    Args:
        all_data: Dictionary mapping region names to DataFrames
        output_path: Path to save the Word document
        source_folder: Path to the source folder (for metadata)
    """
    print(f"\nGenerating report: {output_path}")
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Job Data Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Report metadata
    p = doc.add_paragraph()
    p.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Source folder information
    if source_folder:
        p = doc.add_paragraph()
        p.add_run(f'Data Source: {source_folder}').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Executive Summary
    add_heading(doc, 'Executive Summary', 1)
    
    total_jobs = sum(len(df) for df in all_data.values() if not df.empty)
    regions_with_data = [r for r, df in all_data.items() if not df.empty]
    
    summary_text = f"""
This report analyzes job posting data across {len(regions_with_data)} regions, containing a total of {total_jobs:,} job records.
The analysis covers data completeness, salary distributions, temporal patterns, geographic distribution, and platform sources.
    """
    add_paragraph(doc, summary_text.strip())
    
    # Overall Statistics Table
    add_heading(doc, 'Overall Statistics', 2)
    overall_stats = {
        'Total Regions': len(regions_with_data),
        'Total Jobs': f"{total_jobs:,}",
        'Regions Analyzed': ', '.join(regions_with_data)
    }
    if source_folder:
        overall_stats['Data Source'] = str(source_folder)
    add_table_from_dict(doc, overall_stats)
    
    doc.add_page_break()
    
    # Region-by-Region Analysis
    add_heading(doc, 'Regional Analysis', 1)
    
    for region_name, df in all_data.items():
        if df.empty:
            continue
        
        add_heading(doc, region_name, 2)
        
        # Basic Statistics
        add_heading(doc, 'Basic Statistics', 3)
        basic_stats = {
            'Total Jobs': len(df),
            'Date Range': f"{df['posted_date'].min() if 'posted_date' in df.columns and df['posted_date'].notna().any() else 'N/A'} to {df['posted_date'].max() if 'posted_date' in df.columns and df['posted_date'].notna().any() else 'N/A'}"
        }
        add_table_from_dict(doc, basic_stats)
        
        # Completeness Analysis
        add_heading(doc, 'Data Completeness', 3)
        completeness = calculate_completeness(df)
        if completeness:
            comp_table = doc.add_table(rows=1, cols=3)
            comp_table.style = 'Light Grid Accent 1'
            header_cells = comp_table.rows[0].cells
            header_cells[0].text = 'Field'
            header_cells[1].text = 'Count'
            header_cells[2].text = 'Percentage'
            for cell in header_cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            for field, stats in completeness.items():
                row_cells = comp_table.add_row().cells
                row_cells[0].text = field.replace('_', ' ').title()
                row_cells[1].text = str(stats['count'])
                row_cells[2].text = f"{stats['percentage']:.1f}%"
        
        # Salary Analysis
        add_heading(doc, 'Salary Analysis', 3)
        salary_info = analyze_salary(df, region_name)
        if salary_info.get('has_data'):
            salary_stats = {
                'Sample Size': salary_info['count'],
                'Mean Salary (USD)': f"${salary_info['mean']:,.0f}",
                'Median Salary (USD)': f"${salary_info['median']:,.0f}",
                'Standard Deviation': f"${salary_info['std']:,.0f}",
                'Minimum': f"${salary_info['min']:,.0f}",
                'Maximum': f"${salary_info['max']:,.0f}",
                '25th Percentile': f"${salary_info['q25']:,.0f}",
                '75th Percentile': f"${salary_info['q75']:,.0f}"
            }
            add_table_from_dict(doc, salary_stats)
        else:
            add_paragraph(doc, 'No salary data available for this region.')
        
        # Time Range Analysis
        add_heading(doc, 'Time Range Analysis', 3)
        time_info = analyze_time_range(df)
        if time_info.get('has_data'):
            time_stats = {
                'Jobs with Dates': time_info['count'],
                'Earliest Posting': time_info['earliest'].strftime('%Y-%m-%d'),
                'Latest Posting': time_info['latest'].strftime('%Y-%m-%d'),
                'Time Span': f"{time_info['span_days']} days"
            }
            add_table_from_dict(doc, time_stats)
        else:
            add_paragraph(doc, 'No date information available.')
        
        # Location Analysis
        add_heading(doc, 'Top Locations', 3)
        location_info = analyze_location(df)
        if location_info.get('has_data'):
            top_locs = location_info['top_locations']
            if top_locs:
                loc_table = doc.add_table(rows=1, cols=2)
                loc_table.style = 'Light Grid Accent 1'
                header_cells = loc_table.rows[0].cells
                header_cells[0].text = 'Location'
                header_cells[1].text = 'Count'
                for cell in header_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                
                for loc, count in list(top_locs.items())[:10]:
                    row_cells = loc_table.add_row().cells
                    row_cells[0].text = str(loc)
                    row_cells[1].text = str(count)
        
        # Platform Analysis
        add_heading(doc, 'Platform Distribution', 3)
        platform_info = analyze_platform(df)
        if platform_info.get('has_data'):
            add_table_from_dict(doc, platform_info['distribution'], ['Platform', 'Count'])
        else:
            add_paragraph(doc, 'No platform information available.')
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # Cross-Region Analysis
    add_heading(doc, 'Cross-Region Analysis', 1)
    
    # Create charts directory
    charts_dir = os.path.join(os.path.dirname(output_path), 'report_charts')
    os.makedirs(charts_dir, exist_ok=True)
    
    # Generate and add charts
    chart_paths = {}
    
    chart_path = create_chart_completeness(all_data, charts_dir)
    if chart_path:
        chart_paths['Completeness'] = chart_path
        add_heading(doc, 'Data Completeness Comparison', 2)
        add_paragraph(doc, 'The following chart compares overall data completeness across all regions.')
        doc.add_picture(chart_path, width=Inches(6))
    
    chart_path = create_chart_salary_distribution(all_data, charts_dir)
    if chart_path:
        chart_paths['Salary'] = chart_path
        add_heading(doc, 'Salary Distribution Comparison', 2)
        add_paragraph(doc, 'Salary distributions across regions, showing mean, median, and quartiles.')
        doc.add_picture(chart_path, width=Inches(6))
    
    chart_path = create_chart_time_series(all_data, charts_dir)
    if chart_path:
        chart_paths['Time Series'] = chart_path
        add_heading(doc, 'Job Postings Over Time', 2)
        add_paragraph(doc, 'Monthly job posting trends by region.')
        doc.add_picture(chart_path, width=Inches(6))
    
    chart_path = create_chart_platform_distribution(all_data, charts_dir)
    if chart_path:
        chart_paths['Platform'] = chart_path
        add_heading(doc, 'Platform Distribution', 2)
        add_paragraph(doc, 'Distribution of job postings by platform (Indeed vs LinkedIn).')
        doc.add_picture(chart_path, width=Inches(6))
    
    chart_path = create_chart_region_comparison(all_data, charts_dir)
    if chart_path:
        chart_paths['Comparison'] = chart_path
        add_heading(doc, 'Comprehensive Region Comparison', 2)
        add_paragraph(doc, 'Multi-dimensional comparison of regions across key metrics.')
        doc.add_picture(chart_path, width=Inches(6))
    
    # Cross-Region Statistics Table
    add_heading(doc, 'Cross-Region Statistics', 2)
    
    comparison_data = []
    for region_name, df in all_data.items():
        if df.empty:
            continue
        
        comp = calculate_completeness(df)
        total_fields = len(comp)
        filled_fields = sum(v['count'] for v in comp.values())
        total_records = len(df)
        overall_comp = (filled_fields / (total_fields * total_records) * 100) if total_records > 0 else 0
        
        salary_info = analyze_salary(df, region_name)
        avg_salary = salary_info.get('mean', 0) if salary_info.get('has_data') else 0
        
        comparison_data.append({
            'Region': region_name,
            'Total Jobs': len(df),
            'Completeness (%)': f"{overall_comp:.1f}",
            'Avg Salary (USD)': f"${avg_salary:,.0f}" if avg_salary > 0 else 'N/A'
        })
    
    if comparison_data:
        comp_table = doc.add_table(rows=1, cols=4)
        comp_table.style = 'Light Grid Accent 1'
        header_cells = comp_table.rows[0].cells
        headers = ['Region', 'Total Jobs', 'Completeness (%)', 'Avg Salary (USD)']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        for data in comparison_data:
            row_cells = comp_table.add_row().cells
            row_cells[0].text = data['Region']
            row_cells[1].text = str(data['Total Jobs'])
            row_cells[2].text = data['Completeness (%)']
            row_cells[3].text = data['Avg Salary (USD)']
    
    # Save document
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    
    return output_path


def main():
    """Main function"""
    parser = argparse.ArgumentParser(
        description='Generate comprehensive report from local Excel files',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Analyze a specific folder
  python generate_local_report.py --folder "output/BunchGlobal_2025_01_18"
  
  # Use default folder from configuration
  python generate_local_report.py

Expected folder structure:
  folder_path/
      region1/
          jobspy_max_output.xlsx
      region2/
          jobspy_max_output.xlsx
      ...
        """
    )
    
    parser.add_argument(
        '--folder',
        type=str,
        default=DEFAULT_FOLDER,
        help=f'Path to folder containing region subdirectories with Excel files. Default: {DEFAULT_FOLDER}'
    )
    
    args = parser.parse_args()
    
    if not args.folder:
        print("[ERROR] Please specify a folder path using --folder argument")
        print("Example: python generate_local_report.py --folder \"output/BunchGlobal_2025_01_18\"")
        return
    
    print("="*60)
    print("Local Excel Data Report Generator")
    print("="*60)
    
    try:
        # Load data from folder
        print(f"\nLoading data from folder: {args.folder}")
        all_data = load_data_from_folder(args.folder)
        
        # Check if we have any data
        total_records = sum(len(df) for df in all_data.values() if not df.empty)
        if total_records == 0:
            print("\n[WARNING] No data found in any Excel file. Cannot generate report.")
            return
        
        print(f"\nTotal records loaded: {total_records:,}")
        
        # Generate report
        output_dir = "output/reports"
        os.makedirs(output_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        folder_name = Path(args.folder).name
        output_path = os.path.join(output_dir, f"local_job_analysis_report_{folder_name}_{timestamp}.docx")
        
        generate_report(all_data, output_path, source_folder=args.folder)
        
        print("\n" + "="*60)
        print("Report generation completed!")
        print(f"Report saved to: {output_path}")
        print(f"Data source: {args.folder}")
        print("="*60)
        
    except Exception as e:
        print(f"\n[ERROR] {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()

